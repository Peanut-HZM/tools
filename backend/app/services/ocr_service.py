import base64
import time
import numpy as np
import cv2
import logging
import httpx
import uuid
import json
from typing import List, Dict, Any, Optional
from datetime import datetime, timedelta
from app.models.ocr_models import (
    OCRResponse, TextBlock, QRCodeResponse,
    OCRHistoryRecord, OCRHistoryListResponse,
    ExportOCRRequest, ExportOCRResponse, ExportFormat,
    OCRQuotaInfo, OCRBatchProcessRequest, OCRBatchProcessResponse,
    TableRecognitionRequest, TableRecognitionResponse
)
from app.config.ocr_config import ocr_settings
from app.config.database import get_pooled_db_connection, release_db_connection
from app.core.exceptions import QuotaExceeded
from app.models.base import SessionLocal
from app.services.llm_quota_service import LLMQuotaService

logger = logging.getLogger(__name__)

# Umi-OCR 语言参数映射：短代码 → 模型配置文件路径
LANGUAGE_MAP = {
    "ch": "models/config_chinese.txt",          # 简体中文
    "chinese": "models/config_chinese.txt",
    "en": "models/config_en.txt",               # English
    "english": "models/config_en.txt",
    "cht": "models/config_chinese_cht.txt",     # 繁體中文
    "ja": "models/config_japan.txt",            # 日本語
    "japan": "models/config_japan.txt",
    "ko": "models/config_korean.txt",           # 한국어
    "korean": "models/config_korean.txt",
    "cyrillic": "models/config_cyrillic.txt",   # Русский
}


def _map_language(lang: str) -> str:
    """将短语言代码转换为 Umi-OCR 模型配置路径"""
    return LANGUAGE_MAP.get(lang, lang)


class OCRService:
    def __init__(self):
        self.api_url = ocr_settings.OCR_API_URL
        self.api_key = ocr_settings.API_KEY

    def _ensure_ocr_history_table(self):
        """确保 ocr_history 和 ocr_quota 表存在"""
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ocr_history (
                    id VARCHAR(64) PRIMARY KEY,
                    user_id VARCHAR(64) NOT NULL,
                    original_image_url TEXT,
                    thumbnail_url TEXT,
                    recognized_text TEXT NOT NULL,
                    language VARCHAR(10) DEFAULT 'ch',
                    block_count INT DEFAULT 0,
                    processing_time_ms FLOAT,
                    file_size INT,
                    image_width INT,
                    image_height INT,
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_ocr_history_user_id ON ocr_history(user_id)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_ocr_history_created_at ON ocr_history(created_at)")

            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ocr_quota (
                    user_id VARCHAR(64) PRIMARY KEY,
                    daily_limit INT DEFAULT 100,
                    daily_used INT DEFAULT 0,
                    daily_reset_date DATE,
                    monthly_limit INT DEFAULT 3000,
                    monthly_used INT DEFAULT 0,
                    monthly_reset_date DATE,
                    updated_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE
                )
            """)
            conn.commit()
        finally:
            cursor.close()
            release_db_connection(conn)

    def _decode_image(self, image_data: str) -> np.ndarray:
        try:
            if "base64," in image_data:
                image_data = image_data.split("base64,")[1]
            img_bytes = base64.b64decode(image_data)
            nparr = np.frombuffer(img_bytes, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            return img
        except Exception as e:
            logger.error(f"Image decoding failed: {e}")
            raise ValueError("Invalid image data")

    def _save_history(self, user_id: str, image_data: str, result: OCRResponse, lang: str) -> str:
        """保存 OCR 结果到历史记录"""
        self._ensure_ocr_history_table()
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            history_id = str(uuid.uuid4())

            # 计算图片信息
            img = self._decode_image(image_data) if image_data else None
            file_size = len(base64.b64decode(image_data.split("base64,")[1])) if "base64," in image_data else 0
            img_width, img_height = (img.shape[1], img.shape[0]) if img is not None else (None, None)

            cursor.execute("""
                INSERT INTO ocr_history (
                    id, user_id, original_image_url, thumbnail_url, recognized_text,
                    language, block_count, processing_time_ms, file_size,
                    image_width, image_height, created_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
            """, (
                history_id, user_id, None, None, result.text,
                lang, len(result.blocks), result.processing_time * 1000, file_size,
                img_width, img_height
            ))
            conn.commit()
            return history_id
        except Exception as e:
            conn.rollback()
            logger.error(f"Save OCR history failed: {e}")
            return ""
        finally:
            cursor.close()
            release_db_connection(conn)

    def _check_quota(self, user_id: str) -> OCRQuotaInfo:
        """检查用户配额"""
        self._ensure_ocr_history_table()
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            today = datetime.now().date()
            first_day_of_month = today.replace(day=1)

            # 获取或创建配额记录
            cursor.execute("SELECT * FROM ocr_quota WHERE user_id = %s", (user_id,))
            row = cursor.fetchone()

            if not row:
                # 创建新记录
                cursor.execute("""
                    INSERT INTO ocr_quota (user_id, daily_limit, daily_used, daily_reset_date,
                                          monthly_limit, monthly_used, monthly_reset_date, updated_at)
                    VALUES (%s, 100, 0, %s, 3000, 0, %s, CURRENT_TIMESTAMP)
                """, (user_id, today, today))
                conn.commit()
                return OCRQuotaInfo(
                    user_id=user_id,
                    daily_limit=100,
                    daily_used=0,
                    daily_remaining=100,
                    monthly_limit=3000,
                    monthly_used=0,
                    monthly_remaining=3000,
                    reset_date=datetime.now()
                )

            # 检查是否需要重置每日配额
            daily_reset_date = row['daily_reset_date']
            if daily_reset_date and daily_reset_date < today:
                cursor.execute("""
                    UPDATE ocr_quota SET daily_used = 0, daily_reset_date = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE user_id = %s
                """, (today, user_id))
                row['daily_used'] = 0
                row['daily_reset_date'] = today

            # 检查是否需要重置每月配额
            monthly_reset_date = row['monthly_reset_date']
            if monthly_reset_date and monthly_reset_date < first_day_of_month:
                cursor.execute("""
                    UPDATE ocr_quota SET monthly_used = 0, monthly_reset_date = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE user_id = %s
                """, (first_day_of_month, user_id))
                row['monthly_used'] = 0
                row['monthly_reset_date'] = first_day_of_month

            conn.commit()

            return OCRQuotaInfo(
                user_id=user_id,
                daily_limit=row['daily_limit'],
                daily_used=row['daily_used'],
                daily_remaining=row['daily_limit'] - row['daily_used'],
                monthly_limit=row['monthly_limit'],
                monthly_used=row['monthly_used'],
                monthly_remaining=row['monthly_limit'] - row['monthly_used'],
                reset_date=datetime.combine(row['daily_reset_date'] or today, datetime.min.time())
            )
        except Exception as e:
            logger.error(f"Check quota failed: {e}")
            raise e
        finally:
            cursor.close()
            release_db_connection(conn)

    def _increment_quota_usage(self, user_id: str):
        """增加用户配额使用量"""
        self._ensure_ocr_history_table()
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                UPDATE ocr_quota SET daily_used = daily_used + 1, monthly_used = monthly_used + 1,
                                      updated_at = CURRENT_TIMESTAMP
                WHERE user_id = %s
            """, (user_id,))
            conn.commit()
        except Exception as e:
            conn.rollback()
            logger.error(f"Increment quota failed: {e}")
        finally:
            cursor.close()
            release_db_connection(conn)

    def _quota_session(self):
        """获取 quota 操作用的 SQLAlchemy Session（OCR 自有连接池不暴露 Session，每次创建短生命周期 Session）。"""
        return SessionLocal(), True

    @staticmethod
    def _image_size_kb(image_data: str) -> int:
        """从 base64 编码的图片数据计算图片大小（KB）。"""
        if not image_data:
            return 0
        try:
            payload = image_data.split("base64,", 1)[1] if "base64," in image_data else image_data
            return max(0, len(base64.b64decode(payload)) // 1024)
        except Exception:
            return 0

    def predict(self, image_data: str, lang: str = "ch", user_id: Optional[str] = None, save_history: bool = True) -> OCRResponse:
        """
        调用远程 OCR 服务进行识别
        :param image_data: Base64 编码的图片数据
        :param lang: 语言代码
        :param user_id: 用户 ID（用于配额管理）
        :param save_history: 是否保存历史记录
        """
        start_time = time.time()

        # 检查配额
        if user_id:
            quota = self._check_quota(user_id)
            if quota.daily_remaining <= 0:
                raise ValueError("每日 OCR 次数已用完，请明天再试")
            if quota.monthly_remaining <= 0:
                raise ValueError("每月 OCR 次数已用完，请下月再试")

        # ---- quota 预占（Task 10）：按 image_kb // 4 估算 token ----
        image_kb = self._image_size_kb(image_data)
        planned_tokens = max(1, image_kb // 4)
        quota_db = None
        quota_owns_session = False
        quota_svc = None
        res_id = None
        if user_id:
            quota_db, quota_owns_session = self._quota_session()
            quota_svc = LLMQuotaService(quota_db)
            try:
                res_id = quota_svc.check_and_reserve(
                    user_id=user_id, category="ocr", planned_tokens=planned_tokens,
                )
            except QuotaExceeded as e:
                if quota_owns_session:
                    quota_db.close()
                logger.warning("OCR quota 预占失败: user=%s err=%s", user_id, e)
                raise

        try:
            if "base64," in image_data:
                image_data = image_data.split("base64,")[1]

            payload = {
                "base64": image_data,
                "options": {
                    "ocr.language": _map_language(lang)
                }
            }

            # Umi-OCR 无需认证头，保留 api_key 字段兼容未来扩展
            headers = {}
            if self.api_key:
                headers["api-key"] = self.api_key

            target_url = f"{self.api_url}/api/ocr"
            logger.info(f"Calling OCR API: {target_url}")

            with httpx.Client(timeout=60.0) as client:
                response = client.post(target_url, json=payload, headers=headers)

                if response.status_code != 200:
                    raise RuntimeError(f"OCR API error: {response.status_code} - {response.text}")

                result = response.json()
                code = result.get("code")
                data = result.get("data")
                processing_time = result.get("time", 0.0)

                if code != 100 and code != 101:
                    raise RuntimeError(f"OCR failed with code {code}: {data}")

                blocks = []
                full_text = []

                if isinstance(data, list):
                    for item in data:
                        text = item.get("text", "")
                        confidence = item.get("score", 0.0)
                        box = item.get("box", [])

                        blocks.append(TextBlock(
                            text=text,
                            confidence=confidence,
                            box=box
                        ))
                        full_text.append(text)

                ocr_response = OCRResponse(
                    text="\n".join(full_text),
                    blocks=blocks,
                    processing_time=processing_time
                )

                # ---- quota 校正（Task 10）：用 image_kb // 4 写 usage_log ----
                if quota_svc is not None and res_id is not None:
                    quota_svc.record_usage(
                        user_id=user_id, category="ocr",
                        actual_tokens=planned_tokens, reservation_id=res_id,
                        model_used="umi-ocr",
                    )

                # 保存历史记录和更新配额
                if user_id and save_history:
                    self._save_history(user_id, image_data, ocr_response, lang)
                    self._increment_quota_usage(user_id)

                return ocr_response

        except Exception as e:
            # ---- quota 回滚（Task 10）：OCR 调用失败时回滚预留 ----
            if quota_svc is not None and res_id is not None:
                try:
                    quota_svc.rollback(res_id)
                except Exception:
                    pass
            logger.error(f"Remote OCR prediction failed: {e}")
            raise
        finally:
            if quota_owns_session and quota_db is not None:
                quota_db.close()

    def predict_pdf(self, file_content: bytes, user_id: Optional[str] = None, save_history: bool = True) -> OCRResponse:
        """PDF OCR 识别"""
        try:
            import pypdfium2 as pdfium
        except ImportError:
            raise RuntimeError("pypdfium2 not installed")

        start_time = time.time()
        full_blocks = []
        full_text = []

        # ---- quota 预占（Task 10）：PDF 按文件字节数估算 ----
        pdf_kb = max(1, len(file_content) // 1024)
        planned_tokens = max(1, pdf_kb // 4)
        quota_db = None
        quota_owns_session = False
        quota_svc = None
        res_id = None
        if user_id:
            quota_db, quota_owns_session = self._quota_session()
            quota_svc = LLMQuotaService(quota_db)
            try:
                res_id = quota_svc.check_and_reserve(
                    user_id=user_id, category="ocr", planned_tokens=planned_tokens,
                )
            except QuotaExceeded as e:
                if quota_owns_session:
                    quota_db.close()
                logger.warning("OCR PDF quota 预占失败: user=%s err=%s", user_id, e)
                raise

        try:
            pdf = pdfium.PdfDocument(file_content)
            n_pages = len(pdf)

            with httpx.Client(timeout=60.0) as client:
                target_url = f"{self.api_url}/api/ocr"
                headers = {}
                if self.api_key:
                    headers["api-key"] = self.api_key

                for i in range(n_pages):
                    page = pdf[i]
                    pil_image = page.render(scale=2.0).to_pil()

                    import io
                    buffered = io.BytesIO()
                    pil_image.save(buffered, format="JPEG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

                    payload = {
                        "base64": img_base64,
                        "options": {"ocr.language": _map_language("ch")}
                    }

                    response = client.post(target_url, json=payload, headers=headers)
                    if response.status_code == 200:
                        result = response.json()
                        if result.get("code") in [100, 101]:
                            data = result.get("data", [])
                            if isinstance(data, list):
                                for item in data:
                                    text = item.get("text", "")
                                    confidence = item.get("score", 0.0)
                                    box = item.get("box", [])
                                    full_blocks.append(TextBlock(
                                        text=text,
                                        confidence=confidence,
                                        box=box
                                    ))
                                    full_text.append(text)

            # ---- quota 校正（Task 10） ----
            # 必须在 finally 关闭 quota_db 之前调用 record_usage，
            # 否则会在已关闭的 Session 上写入 llm_usage_log 而失败。
            if quota_svc is not None and res_id is not None:
                try:
                    quota_svc.record_usage(
                        user_id=user_id, category="ocr",
                        actual_tokens=planned_tokens, reservation_id=res_id,
                        model_used="umi-ocr",
                    )
                except Exception:
                    logger.exception("OCR PDF quota 校正失败: user=%s", user_id)

        except Exception as e:
            # ---- quota 回滚（Task 10）：PDF OCR 失败时回滚预留 ----
            if quota_svc is not None and res_id is not None:
                try:
                    quota_svc.rollback(res_id)
                except Exception:
                    pass
            logger.error(f"PDF processing failed: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
        finally:
            if quota_owns_session and quota_db is not None:
                quota_db.close()

        processing_time = time.time() - start_time

        ocr_response = OCRResponse(
            text="\n".join(full_text),
            blocks=full_blocks,
            processing_time=processing_time
        )

        # 保存历史记录（PDF 作为整体保存）
        if user_id and save_history:
            self._save_history(user_id, "", ocr_response, "pdf")
            self._increment_quota_usage(user_id)

        return ocr_response

    def scan_qrcode(self, image_data: str) -> QRCodeResponse:
        """二维码识别"""
        start_time = time.time()
        img = self._decode_image(image_data)

        detector = cv2.QRCodeDetector()
        val, points, straight_qrcode = detector.detectAndDecode(img)

        if not val:
            return QRCodeResponse(
                text="",
                type="None",
                processing_time=time.time() - start_time
            )

        return QRCodeResponse(
            text=val,
            type="QR_CODE",
            processing_time=time.time() - start_time
        )

    def get_history(self, user_id: str, page: int = 1, page_size: int = 20) -> OCRHistoryListResponse:
        """获取用户 OCR 历史记录"""
        self._ensure_ocr_history_table()
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            offset = (page - 1) * page_size

            # 获取总数
            cursor.execute("SELECT COUNT(*) as count FROM ocr_history WHERE user_id = %s", (user_id,))
            total = cursor.fetchone()['count']

            # 获取分页数据
            cursor.execute("""
                SELECT * FROM ocr_history
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            """, (user_id, page_size, offset))

            rows = cursor.fetchall()
            records = [OCRHistoryRecord(
                id=row['id'],
                user_id=row['user_id'],
                original_image_url=row['original_image_url'],
                thumbnail_url=row['thumbnail_url'],
                recognized_text=row['recognized_text'],
                language=row['language'],
                block_count=row['block_count'],
                processing_time_ms=row['processing_time_ms'],
                file_size=row['file_size'],
                image_dimensions={'width': row['image_width'], 'height': row['image_height']} if row['image_width'] else None,
                created_at=row['created_at']
            ) for row in rows]

            return OCRHistoryListResponse(
                records=records,
                total=total,
                page=page,
                page_size=page_size
            )
        except Exception as e:
            logger.error(f"Get OCR history failed: {e}")
            raise e
        finally:
            cursor.close()
            release_db_connection(conn)

    def export_ocr_result(self, user_id: str, request: ExportOCRRequest) -> ExportOCRResponse:
        """导出 OCR 识别结果"""
        self._ensure_ocr_history_table()
        conn = get_pooled_db_connection()
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT * FROM ocr_history WHERE id = %s AND user_id = %s", (request.history_id, user_id))
            row = cursor.fetchone()

            if not row:
                raise ValueError("历史记录不存在")

            text = row['recognized_text']
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

            if request.format == ExportFormat.TXT:
                content = text
                file_name = f"ocr_export_{timestamp}.txt"
            elif request.format == ExportFormat.MD:
                content = f"# OCR 识别结果\n\n{text}"
                file_name = f"ocr_export_{timestamp}.md"
            elif request.format == ExportFormat.JSON:
                content = json.dumps({
                    'text': text,
                    'language': row['language'],
                    'created_at': row['created_at'].isoformat()
                }, ensure_ascii=False, indent=2)
                file_name = f"ocr_export_{timestamp}.json"
            elif request.format == ExportFormat.DOCX:
                # 需要 docx 库
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading('OCR 识别结果', 0)
                    doc.add_paragraph(text)
                    import io
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    content = base64.b64encode(buffer.getvalue()).decode('utf-8')
                    file_name = f"ocr_export_{timestamp}.docx"
                except ImportError:
                    raise ValueError("docx 库未安装")
            else:
                raise ValueError(f"不支持的导出格式：{request.format}")

            return ExportOCRResponse(
                file_name=file_name,
                content=content,
                file_size=len(content.encode('utf-8'))
            )
        except Exception as e:
            logger.error(f"Export OCR result failed: {e}")
            raise e
        finally:
            cursor.close()
            release_db_connection(conn)

    def batch_process(self, user_id: str, request: OCRBatchProcessRequest) -> OCRBatchProcessResponse:
        """批量 OCR 处理"""
        results = []
        history_ids = []
        errors = []
        success_count = 0
        failed_count = 0

        for i, image_data in enumerate(request.images):
            try:
                result = self.predict(image_data, request.lang, user_id, request.auto_save)
                results.append(result)
                if request.auto_save:
                    history_ids.append(self._save_history(user_id, image_data, result, request.lang))
                success_count += 1
            except Exception as e:
                failed_count += 1
                errors.append({'index': i, 'error': str(e)})

        return OCRBatchProcessResponse(
            success_count=success_count,
            failed_count=failed_count,
            results=results,
            history_ids=history_ids,
            errors=errors
        )

    def recognize_table(self, image_data: str, output_format: str = "markdown") -> TableRecognitionResponse:
        """表格识别"""
        start_time = time.time()

        # 先进行普通 OCR 识别
        ocr_result = self.predict(image_data, "ch", save_history=False)

        # 简单的表格检测逻辑（实际应该使用专门的表格识别模型）
        # 这里只是示例，返回原文本
        tables = [ocr_result.text]

        return TableRecognitionResponse(
            tables=tables,
            table_count=len(tables),
            processing_time=time.time() - start_time,
            confidence=0.8  # 示例值
        )

    def get_quota_info(self, user_id: str) -> OCRQuotaInfo:
        """获取用户配额信息"""
        return self._check_quota(user_id)

ocr_service = OCRService()
